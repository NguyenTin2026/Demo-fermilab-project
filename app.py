"""
Fermilab Q&A Assistant — Streamlit front-end.

What this app does, in plain terms:
  1. Loads a fine-tuned Qwen3-4B model (QLoRA adapter or merged weights).
  2. Optionally retrieves Fermilab documents (RAG) to ground the answer.
  3. Lets the user chat, compare RAG vs no-RAG, and view pipeline metrics.

The code is organised as:
  - Setup        : imports, page config, CSS
  - Loaders      : cached functions that load the model, retriever, corpus map, demo cache
  - Helpers      : small reusable functions (generate, build prompt, render UI)
  - UI / Tabs    : the four tabs (Chat, Map, Media, Metrics)

PERFORMANCE NOTES:
  - Inference runs in bf16 on GPU instead of 4-bit nf4. nf4 exists to SAVE VRAM
    during training; at serve time it must dequantize on every matmul, so it is
    slower. Qwen3-4B in bf16 is only ~8GB, which fits easily on an Anvil GPU.
  - The LoRA adapter is folded in with merge_and_unload() to drop per-layer
    adapter overhead.
  - The chat tab streams tokens on GPU (TextIteratorStreamer) so text appears
    progressively. On CPU it falls back to a reliable non-streaming path.
  - If demo_cache.json exists, exact-match demo questions are answered instantly
    from cache (no model call) — ideal for a lag-free live demo.
"""

# ----------------------------------------------------------------------------
# 1. SETUP
# ----------------------------------------------------------------------------
import os
import io                                     # in-memory buffers for PDF/DOCX export bytes
import re                                    # used to "sanitize" downloaded file names (fix #5)
import json
import html as html_lib                     # escape text before putting it in HTML (anti-XSS)
from pathlib import Path
from threading import Lock, Thread           # run model.generate() in the background for streaming
from urllib.parse import urlparse
from datetime import datetime, timezone      # used for chat export (feature 7) and feedback logging (feature 2)

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components  # lets us embed raw HTML/JS in the app
from retriever import Retriever               # our own RAG search module

# Heavy ML deps are OPTIONAL. On Streamlit Community Cloud (CPU, ~1 GB RAM) we do
# NOT load a local LLM -- a 0.5B model in fp32 alone exceeds the memory limit and
# gets the app killed ("Oh no. Error running app."). Instead we answer via the
# Anthropic API (see load_model_and_tokenizer). These imports are guarded so the
# app still boots when torch/transformers/peft are not installed at all.
try:
    import torch
except Exception:
    torch = None
try:
    from transformers import AutoModelForCausalLM, AutoTokenizer, TextIteratorStreamer
except Exception:
    AutoModelForCausalLM = AutoTokenizer = TextIteratorStreamer = None
try:
    from peft import PeftModel, PeftConfig
except Exception:
    PeftModel = PeftConfig = None

# ----------------------------------------------------------------------------
# Optional export dependencies (feature 7 extension: more export formats).
# Both are OPTIONAL: if a package isn't installed, its export format is simply
# left out of the format dropdown instead of crashing the whole app.
#   pip install fpdf2        # enables PDF export
#   pip install python-docx  # enables Word (.docx) export
# (Markdown / plain text / JSON export always work, no extra packages needed.)
# ----------------------------------------------------------------------------
try:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    PDF_EXPORT_AVAILABLE = True
except ImportError:
    PDF_EXPORT_AVAILABLE = False

try:
    from docx import Document
    DOCX_EXPORT_AVAILABLE = True
except ImportError:
    DOCX_EXPORT_AVAILABLE = False

# Optional: lets users attach a PDF as temporary context (feature 16).
#   pip install pypdf
# .txt / .md attachments always work without this.
try:
    from pypdf import PdfReader
    PDF_UPLOAD_AVAILABLE = True
except ImportError:
    PDF_UPLOAD_AVAILABLE = False

# Force HuggingFace to work fully offline (Anvil compute nodes have no internet).
os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")

# How many tokens the model may generate per answer.
# Lower this if you hit GPU out-of-memory; raise it for longer answers.
MAX_NEW_TOKENS = 256

# Total character budget for the RAG context in the prompt (keeps the prompt
# compact to avoid GPU out-of-memory). This budget is now SPLIT EVENLY across
# each retrieved chunk instead of truncating the already-joined string
# (see fix #4 in build_messages).
MAX_CONTEXT_CHARS = 6000

# Number of recent Q&A turns fed into the prompt as conversation context
# (feature 1: conversation memory). 3 turns is enough for the model to
# resolve follow-up questions like "what about X?" without bloating the prompt.
MAX_HISTORY_TURNS = 3

APP_DIR = Path(__file__).parent              # folder this script lives in
FEEDBACK_LOG_PATH = APP_DIR / "feedback_log.jsonl"  # where 👍/👎 feedback is stored (feature 2)

# The fine-tuned model is a single shared GPU resource loaded once (via
# @st.cache_resource below) and reused across EVERY browser session that
# connects to this Streamlit server. If two people click "ask" at nearly the
# same moment, both would otherwise call model.generate() concurrently on the
# same model/GPU state -- transformers' generate() (and the background-thread
# streaming pattern used here) is not safe for that and can corrupt outputs
# or crash the CUDA context. This lock serializes generation across ALL
# sessions: the second request simply waits its turn (still shown as a normal
# spinner to that user) instead of racing the first one.
GENERATION_LOCK = Lock()

st.set_page_config(
    page_title="Fermilab Q&A Assistant Dashboard",
    page_icon="⚛️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Logo Fermilab here!
st.logo(
    "https://noirlab.edu/public/media/archives/logos/screen/logo112.jpg",
    # link="https://www.fnal.gov",
    icon_image="https://noirlab.edu/public/media/archives/logos/screen/logo112.jpg",
)

# Dark "space/physics" theme. Pure decoration — it does not affect any logic.
st.markdown("""
    <style> 
        @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Plus+Jakarta+Sans:wght@300;400;500;600;700&family=Space+Mono:wght@400;700&display=swap');
        html, body, [class*="css"] { font-family: 'Plus Jakarta Sans', sans-serif; }
        .stApp { background-color: #0b0f19; color: #e2e8f0; }
        section[data-testid="stSidebar"] { background-color: #0f172a; border-right: 1px solid rgba(255,255,255,0.05); }
        section[data-testid="stSidebar"] .stMarkdown, section[data-testid="stSidebar"] p { color: #94a3b8; }
        .main-title {
            font-family: 'Outfit', sans-serif; font-weight: 800; font-size: 3rem;
            background: linear-gradient(135deg, #00f2fe 0%, #4facfe 100%);
            -webkit-background-clip: text; -webkit-text-fill-color: transparent;
            margin-bottom: 0.1rem; text-shadow: 0px 4px 12px rgba(0,242,254,0.15);
        }
        .subtitle { font-family: 'Plus Jakarta Sans', sans-serif; font-size: 1.1rem; color: #94a3b8; margin-bottom: 1.5rem; font-weight: 400; }
        .stChatMessage {
            background: rgba(30,41,59,0.4) !important; border: 1px solid rgba(255,255,255,0.05) !important;
            border-radius: 12px !important; margin-bottom: 12px !important; padding: 15px !important; backdrop-filter: blur(10px);
        }
        .source-card {
            background: rgba(30,41,59,0.6); border: 1px solid rgba(0,242,254,0.15); border-left: 4px solid #00f2fe;
            padding: 12px 18px; margin: 10px 0; border-radius: 8px; transition: all 0.3s ease;
        }
        .source-card:hover { transform: translateY(-2px); box-shadow: 0 4px 15px rgba(0,242,254,0.1); border-color: rgba(0,242,254,0.4); }
        .source-title { font-family: 'Outfit', sans-serif; font-weight: 600; color: #38bdf8; font-size: 1rem; margin-bottom: 4px; }
        .source-url { font-family: 'Space Mono', monospace; font-size: 0.8rem; color: #64748b; }
        .source-url a { color: #38bdf8 !important; text-decoration: none; }
        .source-url a:hover { text-decoration: underline; }
        .source-score { font-size: 0.8rem; font-weight: 700; color: #00f2fe; float: right; background: rgba(0,242,254,0.1); padding: 2px 8px; border-radius: 4px; }
        .badge { display: inline-block; padding: 4px 12px; border-radius: 16px; font-size: 0.8rem; font-weight: 700; letter-spacing: 0.05em; margin-bottom: 10px; }
        .badge-rag-on { background-color: rgba(16,185,129,0.15); color: #10b981; border: 1px solid rgba(16,185,129,0.3); }
        .badge-rag-off { background-color: rgba(239,68,68,0.15); color: #ef4444; border: 1px solid rgba(239,68,68,0.3); }
        .pulse-dot { display: inline-block; width: 8px; height: 8px; border-radius: 50%; background-color: #10b981; margin-right: 6px; animation: pulsing 1.5s infinite cubic-bezier(0.66,0,0,1); }
        @keyframes pulsing { to { box-shadow: 0 0 0 8px rgba(16,185,129,0); } }
        .metric-table { width: 100%; border-collapse: collapse; margin: 20px 0; font-size: 0.9rem; background: rgba(30,41,59,0.2); border-radius: 8px; overflow: hidden; border: 1px solid rgba(255,255,255,0.05); }
        .metric-table th, .metric-table td { padding: 12px 15px; text-align: center; border-bottom: 1px solid rgba(255,255,255,0.05); }
        .metric-table th { background-color: rgba(79,172,254,0.1); color: #38bdf8; font-weight: 600; }
    </style>
""", unsafe_allow_html=True)


# ----------------------------------------------------------------------------
# 2. LOADERS (cached so they run only once)
# ----------------------------------------------------------------------------
# Model IDs for the hosted-API backend. Haiku is fast + cheap and is plenty for
# grounded Q&A over short retrieved context -- ideal for a live demo.
API_MODEL = "claude-haiku-4-5-20251001"


def get_api_key():
    """Anthropic API key from Streamlit secrets or the environment (either works)."""
    try:
        key = st.secrets.get("ANTHROPIC_API_KEY", None)
    except Exception:
        key = None
    return key or os.environ.get("ANTHROPIC_API_KEY")


def api_generate(messages, temperature=0.3):
    """Generate an answer with the Anthropic API from build_messages()-style
    messages (a system message + user/assistant turns). Light on memory, so it
    runs fine on Streamlit Cloud."""
    from anthropic import Anthropic

    system = "\n".join(m["content"] for m in messages if m.get("role") == "system")
    chat = [{"role": m["role"], "content": m["content"]}
            for m in messages if m.get("role") in ("user", "assistant")]
    client = Anthropic(api_key=get_api_key())
    resp = client.messages.create(
        model=API_MODEL,
        max_tokens=700,
        temperature=max(0.0, min(1.0, float(temperature))),
        system=system or "You are a helpful assistant.",
        messages=chat or [{"role": "user", "content": "Hello"}],
    )
    return "".join(b.text for b in resp.content if getattr(b, "type", "") == "text").strip()


def grounded_fallback_answer(messages):
    """No API key and no local model: never fabricate. Return the most relevant
    passage from the retrieved context that build_messages() embedded in the user
    message, so the reply is still 100% from the corpus."""
    user = next((m["content"] for m in reversed(messages) if m.get("role") == "user"), "")
    m = re.search(r"Context:\n(.*?)\n\nQuestion:", user, re.S)
    context = (m.group(1) if m else "").strip()
    if not context:
        return ("⚠️ No `ANTHROPIC_API_KEY` is configured, so generated answers are off. "
                "Add the key in **Manage app → Settings → Secrets** to enable full answers. "
                "Retrieved sources are shown below.")
    # first non-header line of the top document
    for line in context.splitlines():
        s = line.strip()
        if s and not s.startswith("---") and not s.lower().startswith(("content:", "url:", "title:")):
            excerpt = s
            break
    else:
        excerpt = context[:600]
    return ("*(No `ANTHROPIC_API_KEY` set — showing the grounded source passage instead of a "
            "generated answer. Add the key in Streamlit secrets for full answers.)*\n\n"
            f"> {excerpt[:700]}")


@st.cache_resource
def load_model_and_tokenizer():
    """
    Pick the answer backend, best-for-Streamlit-Cloud first:
      1. Anthropic API  -- if ANTHROPIC_API_KEY is set (recommended; low memory).
      2. Local model    -- only if torch/transformers are installed AND a GPU is
                           present (dev/SLURM); never attempted on the free CPU tier.
      3. Grounded-only  -- no key, no GPU: return corpus passages, never fabricate.
    Returns: (backend, tokenizer, status_label) where `backend` is the string
    "api"/"grounded" or an actual local model object.
    """
    if get_api_key():
        return "api", None, f"Claude API ({API_MODEL})"

    if torch is None or AutoModelForCausalLM is None or not torch.cuda.is_available():
        # No GPU / no ML stack -> do not load a local model (would OOM on Cloud).
        return "grounded", None, "Retrieval-only (grounded, no LLM)"

    merged_path = APP_DIR / "ft-qwen3-fermilab" / "merged"
    adapter_path = APP_DIR / "ft-qwen3-fermilab"

    # Read the base model name straight from the adapter so it always matches training.
    try:
        base_model_name = PeftConfig.from_pretrained(str(adapter_path)).base_model_name_or_path
    except Exception:
        base_model_name = "Qwen/Qwen2.5-0.5B-Instruct"  # small CPU-friendly fallback (Streamlit Cloud)

    # Pick dtype + device. bf16 on GPU; float32 on CPU (very slow).
    if torch.cuda.is_available():
        device_map = "auto"
        dtype = torch.bfloat16
        dev_label = "GPU/bf16"
    else:
        device_map = "cpu"
        dtype = torch.float32
        dev_label = "CPU/fp32"
        st.sidebar.info(
            "Running on CPU (no GPU). Using a small model (Qwen2.5-0.5B) so it stays "
            "responsive on limited hardware like Streamlit Cloud."
        )

    # Option 1: merged model (fastest).
    if merged_path.exists():
        try:
            model = AutoModelForCausalLM.from_pretrained(
                str(merged_path), torch_dtype=dtype, device_map=device_map
            )
            tokenizer = AutoTokenizer.from_pretrained(str(merged_path))
            model.eval()
            return model, tokenizer, f"Fine-tuned Merged ({dev_label})"
        except Exception as e:
            st.sidebar.error(f"Failed to load merged model: {e}. Trying adapter...")

    # Option 2: base model + LoRA adapter -> merge_and_unload() to drop adapter overhead.
    if adapter_path.exists():
        try:
            base_model = AutoModelForCausalLM.from_pretrained(
                base_model_name, torch_dtype=dtype, device_map=device_map
            )
            model = PeftModel.from_pretrained(base_model, str(adapter_path))
            model = model.merge_and_unload()   # fold LoRA into the base weights
            tokenizer = AutoTokenizer.from_pretrained(base_model_name)
            model.eval()
            return model, tokenizer, f"Fine-tuned LoRA merged ({dev_label})"
        except Exception as e:
            st.sidebar.error(f"Failed to load base + adapter: {e}. Falling back to base model...")

    # Option 3: plain base model.
    model = AutoModelForCausalLM.from_pretrained(
        base_model_name, torch_dtype=dtype, device_map=device_map, low_cpu_mem_usage=True
    )
    tokenizer = AutoTokenizer.from_pretrained(base_model_name)
    model.eval()
    return model, tokenizer, f"Base Model Untrained ({dev_label})"


@st.cache_resource
def load_retriever():
    """Load the RAG retriever from the processed corpus file."""
    path = APP_DIR / "corpus_st7.jsonl"
    if not path.exists():
        st.warning("corpus_st7.jsonl not found. Run the data pipeline first or restore the corpus file.")
        return None
    try:
        return Retriever(str(path))
    except Exception as exc:
        st.warning(f"Unable to initialize retriever: {exc}")
        return None


def normalize_question(text):
    return re.sub(r"\s+", " ", text.strip().lower())


def normalize_train_chunk_id(raw_chunk_id):
    chunk_id = raw_chunk_id.split("#", 1)[0]
    if chunk_id.startswith("st7/"):
        chunk_id = "text_pages/" + chunk_id.split("/", 1)[1]
    return chunk_id


@st.cache_data
def load_demo_cache():
    """
    Load demo_cache.json if present. Pre-generated answers for a fixed set of demo
    questions are served instantly (no model call), keeping a live demo lag-free.
    """
    path = APP_DIR / "demo_cache.json"
    if not path.exists():
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


# Hand-verified source pins for demo questions whose exact wording is NOT a
# verbatim training question, but whose answer is fully grounded in a known
# corpus page (checked against corpus_st7.jsonl). Pinning routes the question
# through the exact-training path: the correct source page is fetched and merged
# as the priority chunk, the grounding guardrail never refuses it, and the model
# answers strictly from that page -- so there is no room to fabricate.
#   NOvA  -> "The NOvA experiment published some of the most precise neutrino
#            oscillation measurements to date ..." (tag/nova page)
#   SeaQuest -> "Fermilab's SeaQuest experiment is hailed for discovering a
#            'sea' of quarks surging inside the proton ..." (tag/quarks page)
DEMO_SOURCE_PINS = {
    "What does Fermilab's NOvA experiment study?": "text_pages/0105_news.fnal.gov_tag_nova",
    "What did the SeaQuest experiment measure?": "text_pages/0725_news.fnal.gov_tag_quarks",
}


@st.cache_data
def load_train_qa_sources():
    """Map exact training questions to the corpus page they were generated from.

    Reads the curated set (train_qa.jsonl) and, when present, the larger
    auto-generated set (train_qa_full.jsonl), then layers the hand-verified demo
    pins on top. Precedence (last write wins): train_qa_full < train_qa < pins.
    Chunk ids are normalized via normalize_train_chunk_id, which strips the
    trailing "#cN" chunk suffix so ids resolve to whole pages in corpus_st7.jsonl.
    """
    sources = {}
    # Widest file first so the curated file and pins can overwrite its entries.
    for fname in ("train_qa_full.jsonl", "train_qa.jsonl"):
        path = APP_DIR / fname
        if not path.exists():
            continue
        try:
            with open(path, "r", encoding="utf-8") as f:
                for raw_line in f:
                    if not raw_line.strip():
                        continue
                    record = json.loads(raw_line)
                    messages = record.get("messages", [])
                    question = next((m.get("content", "") for m in messages if m.get("role") == "user"), "")
                    chunk_id = normalize_train_chunk_id(record.get("meta", {}).get("chunk_id", ""))
                    if question and chunk_id:
                        sources[normalize_question(question)] = {
                            "chunk_id": chunk_id,
                            "origin": record.get("meta", {}).get("origin", ""),
                        }
        except Exception:
            continue
    for question, chunk_id in DEMO_SOURCE_PINS.items():
        sources[normalize_question(question)] = {"chunk_id": chunk_id, "origin": ""}
    return sources


@st.cache_data
def load_corpus_map():
    """
    Return a DataFrame of 2D points (one per document chunk) for the topology map.
    If a cached corpus_map.json exists, reuse it. Otherwise compute it with
    TF-IDF + TruncatedSVD and save it for next time.
    """
    map_path = APP_DIR / "corpus_map.json"
    if map_path.exists():
        try:
            with open(map_path, "r", encoding="utf-8") as f:
                return pd.DataFrame(json.load(f))
        except Exception:
            pass  # unreadable file -> rebuild below

    corpus_path = APP_DIR / "corpus_st7.jsonl"
    if not corpus_path.exists():
        return pd.DataFrame()

    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.decomposition import TruncatedSVD

    # Read every JSONL line into a list of documents.
    documents = []
    with open(corpus_path, "r", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                documents.append(json.loads(line.strip()))
    if not documents:
        return pd.DataFrame()

    # Turn text into vectors, then squash to 2 dimensions for plotting.
    texts = [doc["text"] for doc in documents]
    tfidf = TfidfVectorizer(max_features=1000, stop_words="english").fit_transform(texts)
    coords = TruncatedSVD(n_components=2, random_state=42).fit_transform(tfidf)

    map_data = [
        {
            "x": float(coords[i, 0]),
            "y": float(coords[i, 1]),
            "title": doc.get("title", "Untitled")[:80],
            "source_type": doc.get("source_type", "webpage_text"),
            "split": doc.get("split", "train"),
        }
        for i, doc in enumerate(documents)
    ]

    with open(map_path, "w", encoding="utf-8") as f:
        json.dump(map_data, f, indent=2)
    return pd.DataFrame(map_data)


# Load everything once, with a spinner for the first (slow) run.
with st.spinner("⚛️ Loading Fermilab model and index (first run may take 1-2 minutes)..."):
    model, tokenizer, model_status = load_model_and_tokenizer()
    retriever = load_retriever()
    corpus_map_df = load_corpus_map()
    demo_cache = load_demo_cache()
    train_qa_sources = load_train_qa_sources()


# ----------------------------------------------------------------------------
# 3. HELPERS (small reusable functions)
# ----------------------------------------------------------------------------
# Shared stop-word set for query understanding. Besides ordinary English
# function words it deliberately drops RAG/instruction fluff ("only", "retrieved",
# "documents", "without", "adding", "outside", "knowledge", "context", ...) so a
# question phrased like "using only the retrieved documents without adding outside
# knowledge, what did the Tevatron discover?" reduces to its real content term
# ("tevatron") instead of a pile of generic words the physics corpus never contains.
STOPWORDS = frozenset({
    # articles / pronouns / determiners / prepositions / conjunctions
    "the", "a", "an", "and", "or", "but", "if", "of", "to", "in", "on", "at", "by",
    "for", "with", "as", "is", "are", "was", "were", "be", "been", "being", "am",
    "this", "that", "these", "those", "it", "its", "their", "them", "they", "there",
    "here", "his", "her", "our", "your", "you", "we", "us", "he", "she", "him", "i",
    "what", "which", "who", "whom", "whose", "why", "how", "when", "where", "while",
    "do", "does", "did", "done", "doing", "can", "could", "would", "should", "shall",
    "will", "may", "might", "must", "have", "has", "had", "having", "get", "got",
    "into", "from", "about", "over", "under", "between", "among", "within", "without",
    "through", "during", "before", "after", "above", "below", "than", "then", "up",
    "some", "any", "all", "each", "every", "both", "few", "more", "most", "other",
    "such", "no", "nor", "not", "only", "own", "same", "so", "too", "very", "just",
    "also", "because", "however", "therefore", "thus", "hence", "using", "use",
    "used", "uses", "via", "per", "upon", "out", "off", "again", "further",
    # instruction / meta words that pollute retrieval on RAG-style prompts
    "answer", "answers", "question", "questions", "explain", "explains", "describe",
    "describes", "tell", "list", "name", "give", "provide", "provided", "based",
    "according", "please", "show", "find", "context", "document", "documents",
    "retrieved", "adding", "add", "outside", "knowledge", "info", "information",
    "source", "sources", "summarize", "summary", "detail", "details", "regarding",
    "related", "does", "did",
    # physics-generic words (already present in the old, smaller lists)
    "experiment", "experiments", "study", "studies", "measure", "measured",
    "measurement", "biggest", "discovery", "recently", "primary", "difference",
    "differences", "role", "roles", "contribute", "contributes", "contribution",
    "research", "work", "works", "recent", "new",
})


def _term_covered(term, text):
    """True if `term` appears in `text`, tolerant of simple plural/verb endings
    (so 'discovered' matches 'discovery'/'discoveries', 'proton' matches
    'protons', etc.). Used only to decide whether to refuse -- being lenient here
    just means we answer more often, which is the safe direction."""
    if term in text:
        return True
    if term + "s" in text:
        return True
    for suf in ("s", "es", "ed", "ing"):
        if term.endswith(suf) and len(term) - len(suf) >= 3 and term[: -len(suf)] in text:
            return True
    return False


def select_context_snippet(text, question, budget):
    """Pick the most question-relevant excerpt instead of the first page bytes."""
    if len(text) <= budget:
        return text

    query_terms = [term for term in re.findall(r"\b[a-z0-9]{3,}\b", question.lower()) if term not in STOPWORDS]
    if not query_terms:
        return text[:budget]

    physics_terms = {
        "antiquark", "antiquarks", "collider", "neutrino", "neutrinos",
        "oscillation", "oscillations", "proton", "protons", "quark", "quarks",
        "top", "tevatron", "seaquest", "nova",
    }

    # Score each line and remember the CHARACTER offset of the best-scoring line,
    # then return a window of `budget` characters centered on it. Returning a
    # window (rather than just the few best lines) keeps adjacent evidence inside
    # the snippet -- e.g. the clean English summary sentence that sits right next
    # to a foreign-language reprint of the same result on Fermilab tag pages --
    # so the model sees the quotable answer, not only its noisy neighbour.
    best_score = -1
    best_pos = 0
    pos = 0
    for line in text.splitlines(keepends=True):
        low = line.lower()
        score = sum(low.count(term) for term in query_terms) * 4
        score += sum(1 for term in physics_terms if term in low)
        if score > best_score:
            best_score = score
            best_pos = pos
        pos += len(line)

    if best_score <= 0:
        text_lower = text.lower()
        best_pos = min((text_lower.find(term) for term in query_terms if term in text_lower), default=0)
        best_pos = max(best_pos, 0)

    start = max(0, best_pos - budget // 3)
    return text[start:start + budget]


def important_query_terms(question):
    return [term for term in re.findall(r"\b[a-z0-9]{3,}\b", question.lower()) if term not in STOPWORDS]


def find_train_source_chunk(question, retriever_obj, train_sources):
    record = train_sources.get(normalize_question(question))
    if not record:
        return None, None
    if retriever_obj is None:
        return None, record

    target_id = record["chunk_id"]
    for doc in retriever_obj.documents:
        if doc.get("id") == target_id:
            return {
                "id": doc["id"],
                "text": doc["text"],
                "title": doc["title"],
                "origin": doc["origin"],
                "split": doc.get("split", ""),
                "score": 1.0,
            }, record
    return None, record


def merge_priority_chunk(priority_chunk, chunks, limit):
    merged = []
    seen_ids = set()
    for chunk in [priority_chunk] + list(chunks):
        if not chunk:
            continue
        chunk_id = chunk.get("id")
        if chunk_id in seen_ids:
            continue
        seen_ids.add(chunk_id)
        merged.append(chunk)
        if len(merged) >= limit:
            break
    return merged


def grounding_refusal_reason(question, chunks, is_exact_train_question=False):
    if not chunks:
        return "I don't know from the current Fermilab corpus. No relevant document was retrieved."

    # An exact training question already had its gold source chunk merged into
    # `chunks` (see merge_priority_chunk at the call sites), so we KNOW the
    # evidence is present -- never refuse those on a keyword-coverage heuristic.
    if is_exact_train_question:
        return None

    terms = important_query_terms(question)
    if not terms:
        return None

    # Check coverage against the FULL retrieved text, not a short leading snippet:
    # the deployed corpus stores whole pages, so the relevant sentence often sits
    # well past the first ~1200 characters and must not be judged "missing".
    context = "\n".join(c.get("text", "") for c in chunks).lower()
    covered = [term for term in dict.fromkeys(terms) if _term_covered(term, context)]

    # Refuse only when the retrieved documents share essentially NO content
    # vocabulary with the question -- i.e. retrieval genuinely returned something
    # off-topic. Over-refusing is far more damaging to the demo than occasionally
    # answering from a weakly-matched page, so we err toward answering.
    if not covered:
        missing = ", ".join(dict.fromkeys(terms))
        return (
            "I don't know from the current Fermilab corpus. The retrieved documents do not "
            f"appear to cover this question (no overlap with: {missing})."
        )
    return None


def build_messages(question, chunks, history=None):
    """
    Build the chat messages for the model.
    If `chunks` is non-empty, we run in RAG mode and ground the answer in them.
    Otherwise we ask the model to answer directly from its own knowledge.

    `history` (feature 1): list of PRIOR turns in the current chat session
    (shaped like [{"role": "user"/"assistant", "content": "..."}]).
    It is inserted BETWEEN the system prompt and the current question, so the
    model can understand follow-up questions like "what about that?" instead
    of always treating each question as fully independent.
    """
    if chunks:
        # FIX #4: previously the code joined the text of every chunk first
        # and then sliced context[:6000] -> if the first chunk was long, the
        # last chunk(s) could get truncated mid-sentence (or lose all their
        # content while the "Document [N]" header still appeared).
        # Fix: split the character budget across chunks BEFORE joining, so every
        # document gets a fair chance to contribute to the prompt. The FIRST
        # chunk is the priority/pinned source (an exact training question merges
        # its verified gold page in front via merge_priority_chunk), so it gets
        # the largest share -- a short source page then comes through (nearly)
        # whole and the model sees the full quotable answer, not a fragment --
        # with the remaining budget split evenly across the supporting chunks.
        n_chunks = len(chunks)
        if n_chunks == 1:
            first_budget, rest_budget = MAX_CONTEXT_CHARS, 0
        else:
            first_budget = MAX_CONTEXT_CHARS // 2
            rest_budget = max(1, (MAX_CONTEXT_CHARS - first_budget) // (n_chunks - 1))
        context = ""
        for i, c in enumerate(chunks, start=1):
            budget = first_budget if i == 1 else rest_budget
            snippet = select_context_snippet(c["text"], question, budget)
            context += f"--- Document [{i}]: {c['title']} ---\nContent:\n{snippet}\n\n"

        system_prompt = (
            "You are a helpful assistant that answers questions about Fermilab research.\n"
            "Answer using ONLY the provided context. Every factual claim must be supported "
            "by the context. If the context is missing, weak, or about a different topic, "
            "say you do not know. Do not use outside knowledge, invent facts, or guess beyond the text.\n"
            "If earlier turns in the conversation are provided, use them only to resolve "
            "references (e.g. 'it', 'that experiment') — the factual answer must still come "
            "only from the provided context."
        )
        user_content = f"Context:\n{context}\n\nQuestion: {question}\n\nAnswer:"
    else:
        system_prompt = (
            "You are a helpful assistant that answers questions about Fermilab research. "
            "Answer the question directly and concisely."
        )
        user_content = question

    messages = [{"role": "system", "content": system_prompt}]
    if history:
        messages.extend(history)
    messages.append({"role": "user", "content": user_content})
    return messages


def _build_gen_kwargs(temperature):
    """Sampling config shared by both the streaming and non-streaming paths."""
    gen_kwargs = {
        "max_new_tokens": MAX_NEW_TOKENS,
        "do_sample": temperature > 0.15,   # very low temperature -> deterministic
    }
    if gen_kwargs["do_sample"]:
        gen_kwargs["temperature"] = temperature
        gen_kwargs["top_p"] = 0.9
    return gen_kwargs


def generate_answer(messages, temperature):
    """
    Non-streaming path: returns the full answer at once.
    Routes to the Anthropic API or the grounded fallback when no local model is
    loaded (the Streamlit Cloud case); otherwise uses the local model.
    """
    if model == "api":
        return api_generate(messages, temperature)
    if model == "grounded" or model is None or tokenizer is None:
        return grounded_fallback_answer(messages)

    inputs = tokenizer.apply_chat_template(
        messages,
        add_generation_prompt=True,   # add the marker telling the model to start answering
        return_tensors="pt",
        return_dict=True,
    ).to(model.device)

    with torch.no_grad():                  # inference only, no gradients
        output = model.generate(**inputs, **_build_gen_kwargs(temperature))

    # Keep only the newly generated tokens (drop the prompt part).
    new_tokens = output[0][inputs["input_ids"].shape[-1]:]
    return tokenizer.decode(new_tokens, skip_special_tokens=True).strip()


def generate_answer_stream(messages, temperature):
    """
    Streaming path: yields text chunks as the model produces them.
    model.generate() runs in a background thread while the for-loop reads the streamer.

    IMPORTANT: if generate() raises inside the thread, the error would be swallowed
    and the streamer would end empty (an empty bubble with no message). We capture
    the error and re-raise it after the loop so the caller can see it and fall back.
    """
    # API / grounded backends don't stream token-by-token here -> emit the whole
    # answer once so callers (st.write_stream) still work.
    if model in ("api", "grounded") or model is None or tokenizer is None:
        yield generate_answer(messages, temperature)
        return

    inputs = tokenizer.apply_chat_template(
        messages,
        add_generation_prompt=True,
        return_tensors="pt",
        return_dict=True,
    ).to(model.device)

    streamer = TextIteratorStreamer(
        tokenizer, skip_prompt=True, skip_special_tokens=True, timeout=180
    )

    gen_kwargs = dict(**inputs, streamer=streamer, **_build_gen_kwargs(temperature))

    error_box = {}

    def _worker():
        try:
            with torch.no_grad():
                model.generate(**gen_kwargs)
        except Exception as e:
            error_box["error"] = e   # keep the error to surface it later

    # generate() blocks until done -> run it in a thread so the for-loop can read in parallel.
    thread = Thread(target=_worker)
    thread.start()

    for token in streamer:
        yield token
    thread.join()

    if "error" in error_box:
        raise error_box["error"]   # surface the error so the caller can fall back


def render_source_card(src):
    """Render one retrieved document as a styled card. Text is HTML-escaped (anti-XSS)."""
    title = html_lib.escape(str(src["title"]))
    origin = html_lib.escape(str(src["origin"]))
    st.markdown(f"""
        <div class="source-card">
            <span class="source-score">Relevance: {src['score']}</span>
            <div class="source-title">{title}</div>
            <div class="source-url"><a href="{origin}" target="_blank">{origin}</a></div>
        </div>
    """, unsafe_allow_html=True)


def get_recent_history(exclude_last_n=1):
    """
    (Feature 1) Fetch up to MAX_HISTORY_TURNS recent Q&A turns from
    st.session_state.messages to feed into the prompt as conversation context.

    `exclude_last_n`: number of TRAILING messages to skip before taking history.
      - When handling a new question: exclude_last_n=1 (skip the question we
        just appended, since it's already passed to build_messages separately).
      - When "Regenerate" is used (feature 5): exclude_last_n=2 (skip the
        trailing Q&A pair, since that exact pair is being regenerated).

    Only "role" and "content" are kept — extra fields like "sources" are
    dropped since the model doesn't need them and they'd needlessly bloat
    the prompt.
    """
    msgs = st.session_state.get("messages", [])
    trimmed_end = msgs[:-exclude_last_n] if exclude_last_n > 0 else msgs
    history = [{"role": m["role"], "content": m["content"]} for m in trimmed_end]
    return history[-(MAX_HISTORY_TURNS * 2):]


# ----------------------------------------------------------------------------
# MULTI-SESSION CHAT MANAGEMENT (feature 12) + AUTO-TITLE (feature 11)
# ----------------------------------------------------------------------------
def init_chat_sessions():
    """
    (Feature 12) Set up the multi-session data structure on first run.

    st.session_state.sessions is a dict:
        {session_id: {"title": str, "messages": [ ... ]}}
    and st.session_state.current_session_id points at the active one.

    IMPORTANT TRICK: st.session_state.messages is kept as an ALIAS pointing
    at the CURRENT session's messages list (same list object, not a copy).
    Because Python lists are passed by reference, every existing helper that
    reads or appends to st.session_state.messages keeps working unchanged --
    the mutation lands in the right session automatically. The only rule is:
    NEVER rebind st.session_state.messages to a new list (that would break
    the alias); to clear a chat, call .clear() on the list instead.
    """
    if "sessions" not in st.session_state:
        st.session_state.sessions = {"chat_1": {"title": "New chat", "messages": []}}
        st.session_state.current_session_id = "chat_1"
        st.session_state.session_counter = 1
    # Re-point the alias every run (cheap, and survives session switching).
    st.session_state.messages = (
        st.session_state.sessions[st.session_state.current_session_id]["messages"]
    )


def create_new_session():
    """(Feature 12) Create a fresh chat session and switch to it."""
    st.session_state.session_counter += 1
    new_id = f"chat_{st.session_state.session_counter}"
    st.session_state.sessions[new_id] = {"title": "New chat", "messages": []}
    st.session_state.current_session_id = new_id


def delete_current_session():
    """
    (Feature 12) Delete the active session. Always keeps at least one session
    alive: if the last one is deleted, an empty replacement is created.
    """
    sid = st.session_state.current_session_id
    st.session_state.sessions.pop(sid, None)
    if not st.session_state.sessions:
        st.session_state.session_counter += 1
        new_id = f"chat_{st.session_state.session_counter}"
        st.session_state.sessions[new_id] = {"title": "New chat", "messages": []}
    st.session_state.current_session_id = next(iter(st.session_state.sessions))


def auto_title_session(question):
    """
    (Feature 11) On the FIRST question of a session, name the session after
    that question (trimmed to 40 chars) -- like ChatGPT/Claude auto-titles.
    Later questions never rename the chat, and a manual rename (which changes
    the title away from "New chat") is never overwritten either.
    """
    session = st.session_state.sessions[st.session_state.current_session_id]
    if session["title"] == "New chat" and question.strip():
        title = question.strip()
        session["title"] = title[:40] + ("…" if len(title) > 40 else "")


# ----------------------------------------------------------------------------
# FOLLOW-UP SUGGESTIONS (feature 10)
# ----------------------------------------------------------------------------
def build_followup_suggestions(sources, asked_questions):
    """
    (Feature 10) Build up to 3 suggested follow-up questions from the titles
    of the retrieved source documents -- the idea being that each retrieved
    document is a topic the corpus can actually answer well, so suggesting it
    keeps the user inside well-grounded territory (no extra model call, so
    it's free and instant).

    `asked_questions` filters out suggestions the user has already asked, so
    the same suggestion doesn't reappear turn after turn.
    """
    suggestions, seen_titles = [], set()
    for src in sources or []:
        title = str(src.get("title", "")).strip()
        if not title or title.lower() in seen_titles:
            continue
        seen_titles.add(title.lower())
        question = f'Tell me more about "{title[:60]}"'
        if question.lower() not in asked_questions:
            suggestions.append(question)
        if len(suggestions) == 3:
            break
    return suggestions


def render_followup_suggestions(msg, session_id, idx):
    """
    (Feature 10) Render suggestion buttons under the LATEST answer only.
    Clicking one stores the text in st.session_state.pending_query; the main
    chat handler picks it up on the rerun exactly as if the user had typed it.
    """
    asked = {
        m["content"].strip().lower()
        for m in st.session_state.messages
        if m["role"] == "user"
    }
    suggestions = build_followup_suggestions(msg.get("sources"), asked)
    if not suggestions:
        return
    st.caption("💡 You might also ask:")
    cols = st.columns(len(suggestions))
    for i, (col, sug) in enumerate(zip(cols, suggestions)):
        with col:
            if st.button(sug, key=f"sug_{session_id}_{idx}_{i}", use_container_width=True):
                st.session_state.pending_query = sug
                st.rerun()


# ----------------------------------------------------------------------------
# RETRIEVAL REASONING TRACE (feature 18)
# ----------------------------------------------------------------------------
def render_retrieval_trace(sources, question):
    """
    (Feature 18) A transparent "how did RAG decide?" panel, in the spirit of
    the reasoning traces shown by modern assistants -- but honest about what
    this pipeline actually does: it shows the real retrieval evidence
    (ranked chunks + scores) rather than pretending the LLM "thought".

    Confidence is a simple relative heuristic on the score distribution:
    a clear gap between the top score and the rest usually means one document
    matched much better than everything else (a focused hit); a flat score
    profile means retrieval was less discriminative for this query.
    """
    if not sources:
        return
    scores = []
    for s in sources:
        try:
            scores.append(float(s["score"]))
        except (TypeError, ValueError):
            pass  # e.g. the attached-file pseudo-chunk has a text score

    with st.expander("🔍 Retrieval reasoning (how these sources were chosen)"):
        st.markdown(f"**Query sent to retriever:** `{question}`")
        st.markdown(
            f"**Selection rule:** the retriever scored every chunk in the corpus "
            f"against the query and kept the top {len(sources)} by score."
        )
        for rank, s in enumerate(sources, start=1):
            st.markdown(f"**#{rank}** — {s['title']}  ·  score: `{s['score']}`")
        if len(scores) >= 2:
            top, rest_avg = scores[0], sum(scores[1:]) / len(scores[1:])
            if rest_avg > 0 and top / rest_avg >= 1.5:
                st.info(
                    "📈 Confidence signal: the top document scored clearly higher than "
                    "the rest — retrieval likely found a focused match for this question."
                )
            else:
                st.warning(
                    "📊 Confidence signal: scores are close together — several documents "
                    "are similarly relevant, so the answer may blend multiple sources. "
                    "Consider a more specific question if the answer looks too general."
                )


# ----------------------------------------------------------------------------
# TEMPORARY FILE ATTACHMENT (feature 16)
# ----------------------------------------------------------------------------
# Cap on how much attached-file text goes into the prompt. Kept separate from
# MAX_CONTEXT_CHARS so a huge upload can't crowd out the retrieved corpus.
MAX_UPLOAD_CHARS = 4000


def extract_uploaded_text(uploaded_file):
    """
    (Feature 16) Extract plain text from an uploaded .txt / .md / .pdf file.
    Returns (text, error_message) -- exactly one of the two is None.
    PDF support requires pypdf; without it, PDF uploads get a clear message
    instead of a crash.
    """
    name = uploaded_file.name.lower()
    try:
        if name.endswith((".txt", ".md")):
            return uploaded_file.getvalue().decode("utf-8", errors="replace"), None
        if name.endswith(".pdf"):
            if not PDF_UPLOAD_AVAILABLE:
                return None, "PDF support needs `pip install pypdf`. Try a .txt or .md file instead."
            reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            if not text.strip():
                return None, "No extractable text in this PDF (it may be a scanned image)."
            return text, None
        return None, "Unsupported file type. Please upload .txt, .md, or .pdf."
    except Exception as e:
        return None, f"Could not read the file: {e}"


def get_attached_chunk():
    """
    (Feature 16) If a document is attached to the current session, wrap it as
    one extra pseudo-chunk shaped exactly like a retriever chunk, so
    build_messages() and the source cards can treat it uniformly. The
    attachment is included regardless of the RAG toggle -- attaching a file
    IS an explicit request to ground on it.
    """
    doc = st.session_state.get("attached_doc")
    if not doc:
        return None
    return {
        "title": f"📎 {doc['name']} (attached)",
        "text": doc["text"][:MAX_UPLOAD_CHARS],
        "origin": "user attachment (this session only)",
        "score": "attached",
    }


def _find_preceding_question(msg_index):
    """Find the user question immediately preceding the assistant message at msg_index."""
    msgs = st.session_state.get("messages", [])
    if msg_index > 0 and msgs[msg_index - 1]["role"] == "user":
        return msgs[msg_index - 1]["content"]
    return ""


def log_feedback(question, answer, rating, sources):
    """
    (Feature 2) Append one 👍/👎 feedback line to feedback_log.jsonl.
    This is real signal from actual users — it can later be used to gauge the
    current model's quality or to curate examples for the next fine-tuning
    round. A file-write error (e.g. no write permission on the node) must not
    crash the app — it's only surfaced as a light st.toast warning.
    """
    entry = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "question": question,
        "answer": answer,
        "rating": rating,          # "up" or "down"
        "num_sources": len(sources or []),
    }
    try:
        with open(FEEDBACK_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception as e:
        st.toast(f"⚠️ Could not write feedback: {e}")


def render_assistant_actions(idx, msg, is_last):
    """
    Render the row of action buttons under each assistant reply:
      - 👍 / 👎          : log feedback (feature 2). Each message can only be
                            rated once (buttons disable themselves after a click).
      - 🔄 Regenerate    : shown ONLY on the MOST RECENT reply (feature 5) —
                            regenerates the answer using the current sidebar
                            temperature, always calling the model directly
                            (bypassing the demo cache) since the whole point
                            is to try a different answer.
    """
    # (Feature 12) Keys are scoped by session id: with multiple chats, the
    # message at index 3 of chat_1 and index 3 of chat_2 are different
    # messages -- an unscoped "feedback_3" key would leak ratings (and
    # Streamlit widget-key collisions) across sessions.
    sid = st.session_state.current_session_id
    feedback_key = f"feedback_{sid}_{idx}"
    existing = st.session_state.get(feedback_key)

    if is_last:
        col_up, col_down, col_regen, col_status = st.columns([1, 1, 2, 6])
    else:
        col_up, col_down, col_status = st.columns([1, 1, 8])
        col_regen = None

    with col_up:
        if st.button("👍", key=f"up_{sid}_{idx}", disabled=existing is not None):
            st.session_state[feedback_key] = "up"
            log_feedback(_find_preceding_question(idx), msg["content"], "up", msg.get("sources"))
            st.rerun()
    with col_down:
        if st.button("👎", key=f"down_{sid}_{idx}", disabled=existing is not None):
            st.session_state[feedback_key] = "down"
            log_feedback(_find_preceding_question(idx), msg["content"], "down", msg.get("sources"))
            st.rerun()
    if col_regen is not None:
        with col_regen:
            if st.button("🔄 Regenerate", key=f"regen_{sid}_{idx}", use_container_width=True):
                regenerate_last_answer(idx)
                st.rerun()
    with col_status:
        if existing == "up":
            st.caption("✅ Rated: Helpful")
        elif existing == "down":
            st.caption("📝 Feedback recorded: Not great")


def regenerate_last_answer(idx):
    """
    (Feature 5) Regenerate the answer at position `idx` (always the last
    assistant message). Reuses the question + RAG chunks from the original
    question (stashed in session_state when the question was handled), but
    ALWAYS calls the model directly via generate_answer() — bypassing the
    demo_cache — because if the question matched the cache, "regenerating"
    from cache would just return the exact same old answer, defeating the
    whole purpose of this button.
    """
    # (Feature 12) The stashed question/chunks live INSIDE the session dict,
    # so switching chats can never regenerate with another chat's context.
    sid = st.session_state.current_session_id
    session = st.session_state.sessions[sid]
    query = session.get("last_user_query", "")
    if not query:
        return
    chunks = session.get("last_chunks", [])

    # History up to BEFORE the Q&A pair being regenerated (skip the last
    # 2 messages: the original question + the old answer).
    history = get_recent_history(exclude_last_n=2)
    messages = build_messages(query, chunks, history=history)

    with st.spinner("🔄 Regenerating answer..."):
        try:
            new_answer = generate_answer(messages, temperature)
        except Exception as e:
            st.error(f"⚠️ Could not regenerate the answer: {e}")
            return

    st.session_state.messages[idx]["content"] = new_answer
    # The answer changed -> any previous feedback is no longer valid.
    st.session_state.pop(f"feedback_{sid}_{idx}", None)


def build_chat_export_markdown():
    """
    (Feature 7) Combine the entire current chat history (questions, answers,
    and the source references attached to each answer) into one Markdown
    document, downloadable as a demo/report artifact — similar to the
    comparison report in the RAG vs No-RAG tab, but for the WHOLE chat
    session instead of a single question.
    """
    lines = [
        "# Fermilab Q&A Assistant — Chat Session Export",
        "",
        f"_Exported: {datetime.now(timezone.utc).isoformat()}_",
        f"_Model status: {model_status}_",
        "",
    ]
    for msg in st.session_state.get("messages", []):
        role_label = "🧑 User" if msg["role"] == "user" else "🤖 Assistant"
        lines.append(f"### {role_label}")
        lines.append(msg["content"])
        if msg.get("sources"):
            lines.append("")
            lines.append("**Sources:**")
            for i, src in enumerate(msg["sources"], start=1):
                lines.append(f"{i}. [{src['title']}]({src['origin']}) — Score: {src['score']}")
        lines.append("")
        lines.append("---")
        lines.append("")
    return "\n".join(lines)


def build_chat_export_text():
    """
    (Feature 7 extension) Plain-text version of the chat session export --
    same content as the Markdown export but with no Markdown syntax
    (#, **, [text](url)), for people who just want to paste it into an
    email or a plain document.
    """
    lines = [
        "FERMILAB Q&A ASSISTANT -- CHAT SESSION EXPORT",
        f"Exported: {datetime.now(timezone.utc).isoformat()}",
        f"Model status: {model_status}",
        "=" * 60,
        "",
    ]
    for msg in st.session_state.get("messages", []):
        role_label = "USER" if msg["role"] == "user" else "ASSISTANT"
        lines.append(f"[{role_label}]")
        lines.append(msg["content"])
        if msg.get("sources"):
            lines.append("")
            lines.append("Sources:")
            for i, src in enumerate(msg["sources"], start=1):
                lines.append(f"  {i}. {src['title']} - {src['origin']} (score: {src['score']})")
        lines.append("")
        lines.append("-" * 60)
        lines.append("")
    return "\n".join(lines)


def build_chat_export_json():
    """
    (Feature 7 extension) Structured JSON export of the chat session --
    unlike the text-based formats, this keeps role/content/sources as real
    fields, so it can be fed straight into another script, a spreadsheet, or
    a fine-tuning data pipeline instead of just being read as a document.
    """
    payload = {
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "model_status": model_status,
        "messages": st.session_state.get("messages", []),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _pdf_safe_text(text):
    """
    (Helper for PDF export) The base Helvetica font used by fpdf2 only
    supports Latin-1. Replace common "smart" punctuation (em dash, curly
    quotes, ellipsis...) with plain ASCII first, then drop any character
    that still can't be encoded (e.g. emoji) so PDF generation never
    crashes on real model output.
    """
    replacements = {
        "\u2014": "-", "\u2013": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2026": "...",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def build_chat_export_pdf():
    """
    (Feature 7 extension) Render the chat session as a simple PDF using
    fpdf2. Only ever called when PDF_EXPORT_AVAILABLE is True (checked by
    the caller). Returns raw PDF bytes, ready for st.download_button.

    IMPORTANT fpdf2 gotcha (tested and fixed): by default, multi_cell()
    leaves the cursor's x position at the RIGHT edge of the cell after
    drawing a line, instead of resetting it to the left margin. If the very
    next call is another multi_cell(0, ...) ("use the full remaining
    width"), fpdf2 computes almost zero available width from that stale x
    position and raises "Not enough horizontal space to render a single
    character" -- a hard crash. Every multi_cell() call below explicitly
    passes new_x=XPos.LMARGIN, new_y=YPos.NEXT to force the cursor back to
    the left margin on a new line, which is what "printing a new line of
    text" actually should do.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def _line(txt, style="", size=10):
        """Print one full-width line and always return the cursor to the left margin."""
        pdf.set_font("Helvetica", style, size)
        pdf.multi_cell(0, 6, _pdf_safe_text(txt), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, _pdf_safe_text("Fermilab Q&A Assistant - Chat Session Export"),
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    _line(f"Exported: {datetime.now(timezone.utc).isoformat()}", size=9)
    _line(f"Model status: {model_status}", size=9)
    pdf.ln(4)

    for msg in st.session_state.get("messages", []):
        role_label = "User" if msg["role"] == "user" else "Assistant"
        _line(role_label, style="B", size=11)
        _line(msg["content"], size=10)
        if msg.get("sources"):
            for i, src in enumerate(msg["sources"], start=1):
                _line(f"  [{i}] {src['title']} - {src['origin']} (score: {src['score']})",
                      style="I", size=9)
        pdf.ln(3)

    return bytes(pdf.output())


def build_chat_export_docx():
    """
    (Feature 7 extension) Render the chat session as a .docx file using
    python-docx. Only ever called when DOCX_EXPORT_AVAILABLE is True
    (checked by the caller). Returns raw .docx bytes via an in-memory
    buffer, ready for st.download_button. Unlike the PDF path, python-docx
    handles Unicode natively, so no character sanitizing is needed here.
    """
    doc = Document()
    doc.add_heading("Fermilab Q&A Assistant — Chat Session Export", level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"Exported: {datetime.now(timezone.utc).isoformat()}\n").italic = True
    meta.add_run(f"Model status: {model_status}").italic = True

    for msg in st.session_state.get("messages", []):
        role_label = "User" if msg["role"] == "user" else "Assistant"
        doc.add_heading(role_label, level=2)
        doc.add_paragraph(msg["content"])
        if msg.get("sources"):
            src_heading = doc.add_paragraph()
            src_heading.add_run("Sources:").bold = True
            for i, src in enumerate(msg["sources"], start=1):
                doc.add_paragraph(
                    f"{i}. {src['title']} — {src['origin']} (score: {src['score']})",
                    style="List Bullet",
                )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def get_chat_export(export_format):
    """
    (Feature 7) Dispatch to the right builder for the chosen export format.
    Returns (data, mime_type, file_extension) so the caller can wire up a
    single st.download_button regardless of which format was picked.
    """
    if export_format == "Markdown (.md)":
        return build_chat_export_markdown(), "text/markdown", "md"
    if export_format == "Plain text (.txt)":
        return build_chat_export_text(), "text/plain", "txt"
    if export_format == "JSON (.json)":
        return build_chat_export_json(), "application/json", "json"
    if export_format == "PDF (.pdf)":
        return build_chat_export_pdf(), "application/pdf", "pdf"
    if export_format == "Word (.docx)":
        return (
            build_chat_export_docx(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
    return build_chat_export_markdown(), "text/markdown", "md"  # safe fallback


def answer_toolbar(text):
    """
    (Feature 8 + TTS) Render the per-answer toolbar in ONE self-contained
    iframe: a "Read aloud" button (browser speech engine) and a "Copy"
    button (navigator.clipboard). One iframe instead of two keeps the UI
    compact and halves the component overhead per message.

    ESCAPING NOTE (FIX #6, adapted for this context): the old single-button
    version embedded the text inside an onclick='...' HTML *attribute*, where
    apostrophes had to be HTML-entity-escaped (&#x27;) or they would close the
    attribute early. This version instead stores the text once in a <script>
    block -- and the rules there are DIFFERENT: browsers do NOT decode HTML
    entities inside <script>, so html.escape() here would inject literal
    "&quot;" garbage into the JS source and break it with a syntax error.
    Inside <script>, json.dumps() alone already yields a valid JS string
    literal (quotes/backslashes/newlines escaped); the one remaining hazard
    is the answer text containing the sequence "</script>", which the HTML
    parser would treat as the end of the script block regardless of JS string
    boundaries. Escaping "</" as "<\\/" (a no-op inside a JS string) closes
    that hole.

    The copy path prefers navigator.clipboard (needs a secure context --
    localhost via SSH port forwarding counts as secure) and falls back to a
    hidden <textarea> + execCommand("copy") for plain http:// setups.
    """
    safe_text = json.dumps(text).replace("</", "<\\/")
    html_code = f"""
        <script>
            var answerText = {safe_text};
            function speakAnswer() {{
                window.speechSynthesis.cancel();
                var msg = new SpeechSynthesisUtterance(answerText);
                msg.lang = "en-US";
                window.speechSynthesis.speak(msg);
            }}
            function copyAnswer(btn) {{
                function done() {{
                    var old = btn.innerText;
                    btn.innerText = "✅ Copied!";
                    setTimeout(function() {{ btn.innerText = old; }}, 1500);
                }}
                if (navigator.clipboard && navigator.clipboard.writeText) {{
                    navigator.clipboard.writeText(answerText).then(done);
                }} else {{
                    // Fallback for non-secure (plain http) contexts.
                    var ta = document.createElement("textarea");
                    ta.value = answerText;
                    document.body.appendChild(ta);
                    ta.select();
                    document.execCommand("copy");
                    document.body.removeChild(ta);
                    done();
                }}
            }}
        </script>
        <button onclick="speakAnswer()"
            style="background:#1e293b; color:#38bdf8; border:1px solid rgba(56,189,248,0.4);
                   padding:6px 12px; border-radius:6px; cursor:pointer; font-weight:600; font-size:0.85rem;">
            🔊 Read Aloud
        </button>
        <button onclick="copyAnswer(this)"
            style="background:#1e293b; color:#a5b4fc; border:1px solid rgba(165,180,252,0.4);
                   padding:6px 12px; border-radius:6px; cursor:pointer; font-weight:600; font-size:0.85rem; margin-left:8px;">
            📋 Copy
        </button>
    """
    components.html(html_code, height=45)


def voice_input_widget():
    """
    (Feature 17) A mic button that turns speech into text using the
    browser's built-in Web Speech API (webkitSpeechRecognition). No model,
    no server, no extra Python package: recognition runs in the USER'S
    BROWSER, which has normal internet access even when the Anvil compute
    node serving this app is fully offline (the browser only talks to the
    app through the SSH port-forward tunnel).

    Two-layer design, from best to worst case:
      1. On a result, the widget tries to type the transcript straight into
         Streamlit's chat input box (found in the parent page) and fires a
         React-compatible input event, so the user just hits Enter to send.
         This works because components.html() iframes are same-origin with
         the parent page.
      2. If that injection fails (Streamlit markup changed, browser blocked
         parent access), the transcript is shown inside the widget itself so
         it can be copied manually -- degraded, never broken.

    Browser support note shown in the UI: Chrome/Edge have the API;
    Firefox does not (the widget says so instead of silently failing).
    """
    html_code = """
        <div style="display:flex; align-items:center; gap:10px; font-family:sans-serif;">
            <button id="micBtn"
                style="background:#1e293b; color:#f472b6; border:1px solid rgba(244,114,182,0.4);
                       padding:6px 12px; border-radius:6px; cursor:pointer; font-weight:600; font-size:0.85rem;">
                🎤 Speak your question
            </button>
            <span id="micStatus" style="color:#94a3b8; font-size:0.8rem;"></span>
        </div>
        <div id="transcriptBox" style="display:none; margin-top:6px; padding:6px 10px;
             background:#0f172a; color:#e2e8f0; border-radius:6px; font-size:0.85rem;"></div>
        <script>
            var btn = document.getElementById("micBtn");
            var status = document.getElementById("micStatus");
            var box = document.getElementById("transcriptBox");
            var SR = window.SpeechRecognition || window.webkitSpeechRecognition;

            if (!SR) {
                btn.disabled = true;
                btn.style.opacity = 0.5;
                status.innerText = "Voice input needs Chrome or Edge.";
            } else {
                var rec = new SR();
                rec.lang = "en-US";
                rec.interimResults = false;
                rec.maxAlternatives = 1;

                btn.onclick = function() {
                    status.innerText = "🎙️ Listening...";
                    box.style.display = "none";
                    rec.start();
                };
                rec.onerror = function(e) {
                    status.innerText = "Mic error: " + e.error +
                        (e.error === "not-allowed" ? " (allow microphone access)" : "");
                };
                rec.onresult = function(e) {
                    var transcript = e.results[0][0].transcript;
                    // Layer 1: type the transcript into Streamlit's chat input.
                    try {
                        var ta = window.parent.document.querySelector(
                            'textarea[data-testid="stChatInputTextArea"]'
                        ) || window.parent.document.querySelector(
                            'div[data-testid="stChatInput"] textarea'
                        );
                        if (!ta) throw new Error("chat input not found");
                        // React ignores plain .value writes; use the native
                        // setter then dispatch an input event so Streamlit's
                        // React state actually updates.
                        var setter = Object.getOwnPropertyDescriptor(
                            window.parent.HTMLTextAreaElement.prototype, "value"
                        ).set;
                        setter.call(ta, transcript);
                        ta.dispatchEvent(new Event("input", { bubbles: true }));
                        ta.focus();
                        status.innerText = "✅ Heard it — press Enter to send.";
                    } catch (err) {
                        // Layer 2: show the transcript for manual copy.
                        box.innerText = transcript;
                        box.style.display = "block";
                        status.innerText = "Copy the text below into the chat box:";
                    }
                };
            }
        </script>
    """
    components.html(html_code, height=95)


def build_comparison_export(export_format, custom_query, ans_direct, ans_rag, chunks):
    """
    (Feature 7 extension) Same idea as get_chat_export(), but for the
    single-question RAG vs No-RAG comparison report in Mode B. Returns
    (data, mime_type, file_extension).
    """
    sources = [{"title": c["title"], "origin": c["origin"], "score": c["score"]} for c in chunks]

    if export_format == "Markdown (.md)":
        sources_md = "".join(
            f"{i}. [{s['title']}]({s['origin']}) - Score: {s['score']}\n"
            for i, s in enumerate(sources, start=1)
        ) or "No reference found."
        data = f"""# Fermilab Q&A Hallucination Comparison Report
**Question Analyzed:** "{custom_query}"

---

### ❌ Direct Response (No RAG)
> *Status: unchecked recall — may hallucinate*

{ans_direct}

---

### ✅ Grounded Response (With RAG)
> *Status: verified against the crawled corpus*

{ans_rag}

---

### 📚 References retrieved from corpus:
{sources_md}
"""
        return data, "text/markdown", "md"

    if export_format == "Plain text (.txt)":
        sources_txt = "".join(
            f"  {i}. {s['title']} - {s['origin']} (score: {s['score']})\n"
            for i, s in enumerate(sources, start=1)
        ) or "  No reference found.\n"
        data = (
            "FERMILAB Q&A HALLUCINATION COMPARISON REPORT\n"
            f"Question: {custom_query}\n"
            f"{'=' * 60}\n\n"
            f"[DIRECT RESPONSE - NO RAG]\n{ans_direct}\n\n"
            f"[GROUNDED RESPONSE - WITH RAG]\n{ans_rag}\n\n"
            f"[REFERENCES]\n{sources_txt}"
        )
        return data, "text/plain", "txt"

    if export_format == "JSON (.json)":
        payload = {
            "question": custom_query,
            "direct_answer": ans_direct,
            "rag_answer": ans_rag,
            "sources": sources,
            "exported_at": datetime.now(timezone.utc).isoformat(),
        }
        return json.dumps(payload, ensure_ascii=False, indent=2), "application/json", "json"

    if export_format == "PDF (.pdf)":
        # Same fpdf2 cursor gotcha as build_chat_export_pdf(): every
        # multi_cell() call must pass new_x=XPos.LMARGIN, new_y=YPos.NEXT,
        # otherwise the next line has almost no horizontal space left and
        # fpdf2 raises FPDFException.
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        def _line(txt, style="", size=10):
            pdf.set_font("Helvetica", style, size)
            pdf.multi_cell(0, 6, _pdf_safe_text(txt), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font("Helvetica", "B", 15)
        pdf.multi_cell(0, 8, _pdf_safe_text("Fermilab Q&A Hallucination Comparison Report"),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        _line(f'Question: "{custom_query}"', style="I", size=10)
        pdf.ln(4)

        _line("Direct Response (No RAG)", style="B", size=12)
        _line(ans_direct, size=10)
        pdf.ln(3)

        _line("Grounded Response (With RAG)", style="B", size=12)
        _line(ans_rag, size=10)
        pdf.ln(3)

        if sources:
            _line("References", style="B", size=12)
            for i, s in enumerate(sources, start=1):
                _line(f"{i}. {s['title']} - {s['origin']} (score: {s['score']})", size=9)

        return bytes(pdf.output()), "application/pdf", "pdf"

    if export_format == "Word (.docx)":
        doc = Document()
        doc.add_heading("Fermilab Q&A Hallucination Comparison Report", level=1)
        q_para = doc.add_paragraph()
        q_para.add_run(f'Question: "{custom_query}"').italic = True

        doc.add_heading("Direct Response (No RAG)", level=2)
        doc.add_paragraph(ans_direct)

        doc.add_heading("Grounded Response (With RAG)", level=2)
        doc.add_paragraph(ans_rag)

        if sources:
            doc.add_heading("References", level=2)
            for i, s in enumerate(sources, start=1):
                doc.add_paragraph(
                    f"{i}. {s['title']} — {s['origin']} (score: {s['score']})",
                    style="List Bullet",
                )

        buffer = io.BytesIO()
        doc.save(buffer)
        return (
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )

    # Safe fallback -- should never be hit since options come from the same list.
    return "", "text/plain", "txt"


# ----------------------------------------------------------------------------
# 4. SIDEBAR
# ----------------------------------------------------------------------------
# (Feature 12) Must run before anything touches st.session_state.messages:
# it creates the sessions dict on first run and re-points the .messages
# alias at the currently selected session on every rerun.
init_chat_sessions()

# Text header instead of an external image: external URLs break on offline nodes.
st.sidebar.markdown(
    "<div style='font-family:Outfit,sans-serif; font-size:1.5rem; font-weight:800; "
    "color:#38bdf8; padding:8px 0 4px;'></div>",
    unsafe_allow_html=True,
)

# ---- (Feature 12) Chat session manager ----
st.sidebar.header("🗨️ Chats")
if st.sidebar.button("➕ New Chat", use_container_width=True):
    create_new_session()
    st.rerun()

_session_ids = list(st.session_state.sessions.keys())
_current_sid = st.session_state.current_session_id
_selected_sid = st.sidebar.radio(
    "Your chats",
    options=_session_ids,
    index=_session_ids.index(_current_sid),
    format_func=lambda sid: st.session_state.sessions[sid]["title"],
    label_visibility="collapsed",
)
if _selected_sid != _current_sid:
    st.session_state.current_session_id = _selected_sid
    st.rerun()

with st.sidebar.expander("✏️ Rename / delete this chat"):
    _new_title = st.text_input(
        "Chat title",
        value=st.session_state.sessions[_current_sid]["title"],
        key=f"rename_{_current_sid}",
    )
    if _new_title.strip() and _new_title != st.session_state.sessions[_current_sid]["title"]:
        st.session_state.sessions[_current_sid]["title"] = _new_title.strip()[:60]
        st.rerun()
    if st.button("🗑️ Delete this chat", use_container_width=True):
        delete_current_session()
        st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown(f"**Model Status:** `{model_status}`")
if demo_cache:
    st.sidebar.caption(f"⚡ Demo cache active: {len(demo_cache)} instant answers")
st.sidebar.markdown("---")

st.sidebar.header("🎯 Mode Selection")
app_mode = st.sidebar.selectbox(
    "Choose Interface Mode",
    ["💬 Standard Chat Engine", "⚡ RAG vs No-RAG Comparison"],
    help="Comparison mode shows answers with and without RAG side by side.",
)

st.sidebar.markdown("---")
st.sidebar.header("⚙️ Configuration Settings")

top_k = st.sidebar.slider(
    "📚 Context Passages (Top-K)", min_value=1, max_value=5, value=3,
    help="How many passages to pull from the corpus.",
)
temperature = st.sidebar.slider(
    "🔥 Temperature", min_value=0.1, max_value=1.0, value=0.3, step=0.1,
    help="Higher = more creative but more likely to make things up; lower = more factual.",
)

st.sidebar.markdown("---")
if retriever is not None:
    st.sidebar.metric("🗂️ Indexed Corpus Size", f"{len(retriever.documents):,} Chunks")
else:
    st.sidebar.metric("🗂️ Indexed Corpus Size", "N/A (corpus missing)")

st.sidebar.markdown("""
    <div style='font-size:0.8rem; line-height:1.3; color:#64748b; margin-top:15px;'>
    <strong>System Specifications:</strong><br>
    • Base LLM: Qwen3-4B-Instruct<br>
    • Fine-tuning: QLoRA 4-bit SFT (merged for inference)<br>
    • Inference: bf16 + token streaming<br>
    • Retrieval: see retriever.py<br>
    • Voice: HTML5 Speech API
    </div>
""", unsafe_allow_html=True)


# ----------------------------------------------------------------------------
# 5. MAIN TITLE + TABS
# ----------------------------------------------------------------------------
st.markdown("<div class='main-title'>Fermilab Science Q&A Assistant</div>", unsafe_allow_html=True)
st.markdown("<div class='subtitle'>Domain-specific LLM with Retrieval-Augmented Grounding</div>", unsafe_allow_html=True)

tab_chat, tab_media, tab_metrics = st.tabs([
    "💬 Chatbot Portal",
    "🎥 Fermilab Media Center",
    "📊 Crawler & Pipeline Metrics",
])


# ----------------------------------------------------------------------------
# TAB 1 — CHATBOT
# ----------------------------------------------------------------------------
with tab_chat:

    # ---- Mode A: standard chat ----
    if app_mode == "💬 Standard Chat Engine":

        use_rag = True
        st.markdown('<div class="badge badge-rag-on"><span class="pulse-dot"></span>RAG ACTIVE: Grounded Mode</div>', unsafe_allow_html=True)
        st.caption("Standard Chat is locked to corpus-grounded answers. If the corpus does not support an answer, the assistant will say it does not know.")

        # Chat history: st.session_state.messages is an ALIAS to the current
        # session's list (set up by init_chat_sessions). NEVER rebind it with
        # "= []" -- that would silently detach the alias from the session and
        # new messages would vanish on the next rerun. Mutate in place instead.
        if st.button("🧹 Clear Chat History"):
            st.session_state.messages.clear()   # in-place: alias-safe
            st.session_state.sessions[st.session_state.current_session_id]["title"] = "New chat"
            st.rerun()

        # ---- (Feature 16) Attach a temporary document as extra context ----
        attach_types = ["txt", "md"] + (["pdf"] if PDF_UPLOAD_AVAILABLE else [])
        with st.expander("📎 Attach a document (temporary context for this session)"):
            st.caption(
                "The attached file is added to the model's context for every question "
                "in this session — on top of (or instead of) corpus retrieval. It is "
                "never saved to the corpus and disappears when the app restarts."
                + ("" if PDF_UPLOAD_AVAILABLE else "  \n💡 PDF support: `pip install pypdf`")
            )
            uploaded = st.file_uploader(
                "Upload .txt / .md" + (" / .pdf" if PDF_UPLOAD_AVAILABLE else ""),
                type=attach_types, key="attach_uploader",
            )
            if uploaded is not None:
                text, err = extract_uploaded_text(uploaded)
                if err:
                    st.error(err)
                else:
                    st.session_state.attached_doc = {"name": uploaded.name, "text": text}
                    st.success(
                        f"Attached **{uploaded.name}** "
                        f"({len(text):,} chars; first {MAX_UPLOAD_CHARS:,} used per question)."
                    )
            if st.session_state.get("attached_doc"):
                st.info(f"📎 Currently attached: **{st.session_state.attached_doc['name']}**")
                if st.button("❌ Remove attachment"):
                    st.session_state.pop("attached_doc", None)
                    st.rerun()

        # ---- (Feature 17) Voice input (browser speech-to-text) ----
        voice_input_widget()

        # (Feature 7) Export the current chat session in the chosen format.
        # Markdown / Plain text / JSON always work (no extra packages).
        # PDF / Word only appear once fpdf2 / python-docx are installed.
        chat_export_options = ["Markdown (.md)", "Plain text (.txt)", "JSON (.json)"]
        if PDF_EXPORT_AVAILABLE:
            chat_export_options.append("PDF (.pdf)")
        if DOCX_EXPORT_AVAILABLE:
            chat_export_options.append("Word (.docx)")

        col_fmt, col_dl = st.columns([1, 2])
        with col_fmt:
            chat_export_format = st.selectbox(
                "Export format", chat_export_options,
                key="chat_export_format", label_visibility="collapsed",
            )
        with col_dl:
            chat_data, chat_mime, chat_ext = get_chat_export(chat_export_format)
            st.download_button(
                f"📥 Export chat session (.{chat_ext})",
                data=chat_data,
                file_name=f"Fermilab_Chat_Session_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.{chat_ext}",
                mime=chat_mime,
                disabled=len(st.session_state.messages) == 0,
                use_container_width=True,
            )

        if not (PDF_EXPORT_AVAILABLE and DOCX_EXPORT_AVAILABLE):
            missing = []
            if not PDF_EXPORT_AVAILABLE:
                missing.append("`pip install fpdf2`")
            if not DOCX_EXPORT_AVAILABLE:
                missing.append("`pip install python-docx`")
            st.caption(f"💡 More export formats available after: {' and '.join(missing)}")

        # Re-draw the existing conversation.
        num_messages = len(st.session_state.messages)
        for idx, msg in enumerate(st.session_state.messages):
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                if msg["role"] == "assistant":
                    answer_toolbar(msg["content"])
                    if msg.get("sources"):
                        with st.expander("📚 Sources & References Used"):
                            for src in msg["sources"]:
                                render_source_card(src)
                    # (Features 2 & 5) feedback buttons on every reply,
                    # "Regenerate" button only on the most recent reply.
                    render_assistant_actions(idx, msg, is_last=(idx == num_messages - 1))

        # Handle a new question.
        if user_query := st.chat_input("Ask a question about Fermilab research..."):
            # 1. Show + store the user's message.
            with st.chat_message("user"):
                st.markdown(user_query)
            st.session_state.messages.append({"role": "user", "content": user_query})

            # 2. Standard Chat is always grounded, so do not let a stale demo
            # cache bypass the current corpus and retrieval guardrails.
            cache_hit = None
            train_source_chunk, train_source_record = find_train_source_chunk(user_query, retriever, train_qa_sources)
            train_source_missing = train_source_record is not None and train_source_chunk is None

            # 3. Retrieve context if RAG is on (skipped when we already have a cache hit).
            chunks = []
            if cache_hit is None and use_rag:
                if retriever is not None:
                    with st.spinner("🔍 Searching Fermilab corpus..."):
                        chunks = retriever.retrieve(user_query, top_k=top_k)
                        if train_source_chunk is not None:
                            chunks = merge_priority_chunk(train_source_chunk, chunks, top_k)
                else:
                    st.warning("Retriever unavailable: corpus file is missing.")

            grounding_refusal = None
            if train_source_missing:
                grounding_refusal = (
                    "I don't know from the current Fermilab corpus. This exact training question "
                    f"points to source `{train_source_record['chunk_id']}`, but that source is not present in `corpus_st7.jsonl`."
                )
            elif use_rag:
                grounding_refusal = grounding_refusal_reason(
                    user_query, chunks, is_exact_train_question=train_source_record is not None
                )

            # 4. Produce the answer:
            #    - cache hit -> instant
            #    - GPU       -> stream tokens
            #    - CPU       -> reliable non-stream with a clear spinner
            with st.chat_message("assistant"):
                if cache_hit is not None:
                    # FIX #3: use .get() instead of cache_hit["answer"] so a
                    # malformed entry in demo_cache.json doesn't crash the whole app.
                    answer = cache_hit.get("answer", "")
                    # FIX #1: only show cached RAG sources when RAG is ON, so
                    # the RAG ON/OFF badge always matches what's actually shown
                    # (previously sources were shown even with the toggle off).
                    chunks = cache_hit.get("sources", []) if use_rag else []
                    st.markdown(answer)
                elif grounding_refusal:
                    answer = grounding_refusal
                    st.warning(answer)
                else:
                    # (Feature 1) Grab a few recent turns so the model can
                    # understand follow-up questions ("what about X?"). Skip
                    # the last message (exclude_last_n=1) since that's the
                    # current question, already passed to build_messages separately.
                    history = get_recent_history(exclude_last_n=1)
                    messages = build_messages(user_query, chunks, history=history)
                    use_stream = bool(torch is not None and torch.cuda.is_available())
                    # FIX #2: use a dedicated placeholder for the streamed
                    # answer. If generate() errors mid-stream, we clear the
                    # partial text in this placeholder BEFORE drawing the
                    # fallback answer — avoiding two overlapping answers on screen.
                    stream_placeholder = st.empty()
                    try:
                        if use_stream:
                            # [GPU] progressively render tokens into the dedicated placeholder.
                            with stream_placeholder.container():
                                answer = st.write_stream(generate_answer_stream(messages, temperature))
                        else:
                            # [CPU] non-stream: wait until done, then show the full answer at once.
                            with st.spinner("⚛️ Generating on CPU — this can take a few minutes, please wait..."):
                                answer = generate_answer(messages, temperature)
                            stream_placeholder.markdown(answer)
                    except Exception as e:
                        # Clear the partial/broken answer (if any) before
                        # falling back to the safe, non-stream path.
                        stream_placeholder.empty()
                        st.error(f"⚠️ Generation error → retrying in standard mode. Details: {e}")
                        with st.spinner("⚛️ Retrying in standard mode..."):
                            answer = generate_answer(messages, temperature)
                        st.markdown(answer)

                answer_toolbar(answer)

                sources = [
                    {"title": c["title"], "origin": c["origin"], "score": c["score"]}
                    for c in chunks
                ]
                if sources:
                    with st.expander("📚 Sources & References Used"):
                        for src in sources:
                            render_source_card(src)

                # (Feature 5) Save the question + chunks used so the
                # "Regenerate" button can reuse them without re-querying.
                st.session_state.last_user_query = user_query
                st.session_state.last_chunks = chunks
                st.session_state.last_use_rag = use_rag

                # Store the assistant message BEFORE rendering the feedback
                # buttons, so render_assistant_actions can reference the
                # correct index (idx) of the reply we just created.
                st.session_state.messages.append(
                    {"role": "assistant", "content": answer, "sources": sources}
                )
                render_assistant_actions(
                    len(st.session_state.messages) - 1,
                    st.session_state.messages[-1],
                    is_last=True,
                )

    # ---- Mode B: RAG vs No-RAG comparison ----
    else:
        st.markdown("### ⚡ Live Verification: Hallucination Reduction Test")
        st.markdown("Ask about a Fermilab experiment. The model runs twice to show how RAG reduces made-up answers.")

        st.markdown("**💡 Quick demo questions (click to ask):**")
        col_a, col_b, col_c = st.columns(3)

        if "selected_query" not in st.session_state:
            st.session_state.selected_query = "What did the SeaQuest experiment measure?"

        with col_a:
            if st.button("🧬 What did SeaQuest measure?", use_container_width=True):
                st.session_state.selected_query = "What did the SeaQuest experiment measure?"
        with col_b:
            if st.button("🔬 What is the goal of DUNE?", use_container_width=True):
                st.session_state.selected_query = "What is the goal of the DUNE experiment at LBNF?"
        with col_c:
            if st.button("⚛️ What does NOvA study?", use_container_width=True):
                st.session_state.selected_query = "What does Fermilab's NOvA experiment study?"

        custom_query = st.text_input("📝 Edit your question here:", value=st.session_state.selected_query)

        if st.button("🚀 Run Comparative Analysis", type="primary", use_container_width=True):
            # Retrieve context once; the RAG column reuses it.
            chunks = retriever.retrieve(custom_query, top_k=top_k) if retriever is not None else []
            train_source_chunk, train_source_record = find_train_source_chunk(custom_query, retriever, train_qa_sources)
            train_source_missing = train_source_record is not None and train_source_chunk is None
            if train_source_chunk is not None:
                chunks = merge_priority_chunk(train_source_chunk, chunks, top_k)
            rag_refusal = None
            if train_source_missing:
                rag_refusal = (
                    "I don't know from the current Fermilab corpus. This exact training question "
                    f"points to source `{train_source_record['chunk_id']}`, but that source is not present in `corpus_st7.jsonl`."
                )
            else:
                rag_refusal = grounding_refusal_reason(
                    custom_query, chunks, is_exact_train_question=train_source_record is not None
                )

            col_direct, col_rag = st.columns(2)

            # Column 1: direct answer (NO RAG).
            with col_direct:
                st.markdown("""
                    <div style='background:rgba(239,68,68,0.05); padding:15px; border-radius:8px; border:1px solid rgba(239,68,68,0.15); margin-bottom:15px;'>
                        <span style='color:#ef4444; font-weight:700; font-size:1.1rem;'>❌ Generative Model (Direct)</span><br>
                        <span style='color:#94a3b8; font-size:0.85rem;'>Answers from model weights only. Can be factually wrong.</span>
                    </div>
                """, unsafe_allow_html=True)
                with st.spinner("Generating direct response..."):
                    ans_direct = generate_answer(build_messages(custom_query, []), temperature)
                st.markdown("### Answer:")
                st.info(ans_direct)
                answer_toolbar(ans_direct)

            # Column 2: grounded answer (WITH RAG).
            with col_rag:
                st.markdown("""
                    <div style='background:rgba(16,185,129,0.05); padding:15px; border-radius:8px; border:1px solid rgba(16,185,129,0.15); margin-bottom:15px;'>
                        <span style='color:#10b981; font-weight:700; font-size:1.1rem;'>✅ Grounded RAG Model</span><br>
                        <span style='color:#94a3b8; font-size:0.85rem;'>Answers only from retrieved context. Verifiable.</span>
                    </div>
                """, unsafe_allow_html=True)
                with st.spinner("Retrieving facts & generating response..."):
                    if rag_refusal:
                        ans_rag = rag_refusal
                    elif chunks:
                        ans_rag = generate_answer(build_messages(custom_query, chunks), temperature)
                    else:
                        ans_rag = "No relevant document found in the corpus."

                st.markdown("### Answer:")
                if rag_refusal:
                    st.warning(ans_rag)
                elif chunks:
                    st.success(ans_rag)
                    answer_toolbar(ans_rag)
                    st.markdown("---")
                    st.markdown("##### 📚 Reference Sources Grounded from Corpus:")
                    for src in chunks:
                        render_source_card(src)
                else:
                    st.warning(ans_rag)

            # Downloadable comparison report, now in multiple formats.
            st.markdown("---")
            st.markdown("### 📥 Download Session Comparison Report")

            comparison_export_options = ["Markdown (.md)", "Plain text (.txt)", "JSON (.json)"]
            if PDF_EXPORT_AVAILABLE:
                comparison_export_options.append("PDF (.pdf)")
            if DOCX_EXPORT_AVAILABLE:
                comparison_export_options.append("Word (.docx)")

            comparison_format = st.selectbox(
                "Export format", comparison_export_options, key="comparison_export_format",
            )
            comp_data, comp_mime, comp_ext = build_comparison_export(
                comparison_format, custom_query, ans_direct, ans_rag, chunks
            )
            safe_slug = re.sub(r"[^A-Za-z0-9_-]", "", custom_query[:20].replace(" ", "_"))
            st.download_button(
                label=f"📥 Export Comparison Report (.{comp_ext})",
                data=comp_data,
                file_name=f"Fermilab_Comparison_{safe_slug or 'query'}.{comp_ext}",
                mime=comp_mime,
                use_container_width=True,
            )
            if not (PDF_EXPORT_AVAILABLE and DOCX_EXPORT_AVAILABLE):
                missing = []
                if not PDF_EXPORT_AVAILABLE:
                    missing.append("`pip install fpdf2`")
                if not DOCX_EXPORT_AVAILABLE:
                    missing.append("`pip install python-docx`")
                st.caption(f"💡 More export formats available after: {' and '.join(missing)}")



# ----------------------------------------------------------------------------
# TAB 2 — MEDIA
# ----------------------------------------------------------------------------

with tab_media:
    st.markdown("### 🎥 Fermilab Educational & Science Center")
    st.markdown("Explore particle physics and accelerators through video and audio.")

    media_col1, media_col2 = st.columns(2)
    with media_col1:
        st.markdown("#### 📺 Video: What is a Neutrino?")
        st.markdown("A Fermilab educational video about neutrinos.")
        st.video("https://youtu.be/UY1QQr-PZOg?si=ac6WgmfcOQFT1Zz4")
    with media_col2:
        st.markdown("#### 🔊 Audio: Acoustic Sounds of the Universe")
        st.markdown("Sun sound waves recorded by NASA.")
        st.audio("https://www.nasa.gov/wp-content/uploads/2018/07/sun_sonification.wav")
        st.markdown(
            "<p style='font-size:0.8rem; color:#64748b;'>Source: NASA Heliophysics Archive.</p>",
            unsafe_allow_html=True,
        )

    # Link Fermilab đúng cách — chọn 1 trong 2:
    # st.markdown("🔗 Learn more at [Fermilab](https://www.fnal.gov)")   # link chữ trong markdown
    st.link_button("Click here: Fermilab.gov", "https://www.fnal.gov")     # hoặc nút bấm (Streamlit ≥ 1.34)

# ----------------------------------------------------------------------------
# TAB 3 — PIPELINE METRICS
# ----------------------------------------------------------------------------
with tab_metrics:
    st.markdown("### 📈 Crawler Monitor & Pipeline Training Log")

    st.markdown("#### 🕷️ 1. Web Crawler Activity Dashboard")
    report_csv = APP_DIR.parent / "crawled_data" / "download_report_1000.csv"

    if report_csv.exists():
        try:
            df_crawl = pd.read_csv(str(report_csv))

            c1, c2, c3, c4 = st.columns(4)
            total = len(df_crawl)
            success_count = len(df_crawl[df_crawl["status"] == "SUCCESS"])
            with c1:
                st.metric("Total Pages Attempted", total)
            with c2:
                st.metric("Downloaded Successfully", success_count)
            with c3:
                rate = (success_count / total * 100) if total else 0
                st.metric("Crawl Success Rate", f"{rate:.1f}%")
            with c4:
                st.metric("Average Word Count", f"{int(df_crawl['word_count'].mean())} words/page")

            chart_c1, chart_c2 = st.columns(2)
            with chart_c1:
                st.markdown("**Crawled Pages by Domain**")
                df_crawl["domain"] = df_crawl["url"].apply(lambda u: urlparse(str(u)).hostname or "Unknown")
                domain_counts = df_crawl["domain"].value_counts().reset_index()
                domain_counts.columns = ["Domain", "Pages"]
                st.bar_chart(domain_counts, x="Domain", y="Pages")
            with chart_c2:
                st.markdown("**HTTP Status Codes Distribution**")
                status_dist = df_crawl["http_status_code"].value_counts().reset_index()
                status_dist.columns = ["HTTP Status", "Occurrences"]
                status_dist["HTTP Status"] = status_dist["HTTP Status"].apply(
                    lambda x: str(int(x)) if pd.notna(x) else "Timeout/Error"
                )
                st.bar_chart(status_dist, x="HTTP Status", y="Occurrences")
        except Exception as e:
            st.error(f"Failed to read crawl report: {e}")
    else:
        st.warning(f"Report file not found at: {report_csv}. Run the crawler first.")

    st.markdown("---")
    st.markdown("#### 🤖 2. Stage 3: LoRA Training History Log")
    st.markdown("""
        <table class="metric-table">
            <thead>
                <tr>
                    <th>Global Step</th><th>Epoch</th><th>Training Loss</th>
                    <th>Token Acc.</th><th>Val Loss</th><th>Val Token Acc.</th>
                </tr>
            </thead>
            <tbody>
                <tr><td>5</td><td>0.35</td><td>2.7476</td><td>57.69%</td><td>—</td><td>—</td></tr>
                <tr><td>10</td><td>0.70</td><td>1.5354</td><td>67.66%</td><td>—</td><td>—</td></tr>
                <tr><td>15</td><td>1.00</td><td>1.4085</td><td>70.00%</td><td>—</td><td>—</td></tr>
                <tr><td>20</td><td>1.35</td><td>1.1682</td><td>73.37%</td><td>—</td><td>—</td></tr>
                <tr><td>25</td><td>1.70</td><td>1.1433</td><td>73.73%</td><td>1.3397</td><td>70.71%</td></tr>
                <tr style="background-color: rgba(16,185,129,0.1); font-weight: bold;">
                    <td>30</td><td>2.00</td><td>1.1005</td><td>73.88%</td><td>1.3369</td><td>70.64%</td>
                </tr>
            </tbody>
        </table>
    """, unsafe_allow_html=True)

    st.success("🎓 Training complete. Loss 2.748 → 1.101, token accuracy 57.69% → 73.88%. Validation stable (val loss 1.34, val acc 70.64%).")

# Run this file: streamlit run app.py --server.port 8501 --server.address 0.0.0.0
